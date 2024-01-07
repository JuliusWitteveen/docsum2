# -------------------------------------------------------------------
# file_handler.py
# -------------------------------------------------------------------
"""
The file_handler module provides functionality for loading and saving documents. It supports multiple file formats, including PDF, DOCX, RTF, and TXT. This module interacts with external modules such as PyMuPDF, striprtf, and reportlab for handling specific file formats. It also uses the 'config' module for configuration values like supported file formats and summary save paths.
"""

import os
import re
from docx import Document
import fitz  # PyMuPDF
from striprtf.striprtf import rtf_to_text
import logging
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib.pagesizes import letter
import config  # Import the config module

# Configure logging
# Logging is used extensively for debugging and tracking the module's execution,
# especially in error scenarios and during critical operations like file loading and saving.
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# Use values from config module
# 'SUPPORTED_FILE_FORMATS' and 'SUMMARY_SAVE_PATH' are configuration values imported from the 'config' module.
# These are used to determine the formats that can be loaded and saved, and the default path for saving summaries.
SUPPORTED_FILE_FORMATS = config.SUPPORTED_FILE_FORMATS
SUMMARY_SAVE_PATH = config.SUMMARY_SAVE_PATH

def is_valid_file_path(path):
    """
    Validates the file path using a regular expression and checks if the file exists.

    This function is used internally in the 'load_document' function to ensure the file path is valid
    and the file exists before attempting to load it.

    Args:
        path (str): The file path to validate.

    Returns:
        bool: True if the path is valid and the file exists, False otherwise.
    """
    pattern = r'^[a-zA-Z0-9_\\-\\/:. ]+$'
    if not re.match(pattern, path):
        logging.warning(f"Invalid file path format: {path}")
        return False
    if not os.path.isfile(path):
        logging.warning(f"File does not exist: {path}")
        return False
    return True

def load_document(file_path):
    """
    Loads and returns the text content of a document. It supports PDF, DOCX, RTF, and TXT files.

    This function is used externally, primarily in the 'main.py' module for loading the document to be summarized.
    It interacts with 'is_valid_file_path' for file path validation and different libraries for handling various file formats.

    Args:
        file_path (str): The path of the file to load.

    Returns:
        str: The text content of the document.

    Raises:
        ValueError: If the file path is invalid or the file extension is unsupported.

    Exceptions are logged, and a ValueError is raised for unsupported file formats.
    """
    if not is_valid_file_path(file_path):
        raise ValueError(f"Invalid file path: {file_path}")

    file_extension = os.path.splitext(file_path)[1].lower()
    text = ""

    try:
        if file_extension == ".pdf":
            with fitz.open(file_path) as doc:
                text = "\n".join(page.get_text() for page in doc)
        elif file_extension == ".docx":
            doc = Document(file_path)
            text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
        elif file_extension == ".rtf":
            with open(file_path, 'r', encoding='utf-8') as file:
                rtf_text = file.read()
            text = rtf_to_text(rtf_text)
        elif file_extension == ".txt":
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
        else:
            raise ValueError(f"Unsupported file extension: {file_extension}")
    except Exception as e:
        logging.error(f"Error loading document from {file_path}: {e}")
        raise

    logging.info(f"Document loaded successfully from {file_path}")
    return text

def save_summary(summary, file_path):
    """
    Saves the summary to a file.

    Args:
        summary (str): The summary text to save.
        file_path (str): The path where the summary should be saved.
    """
    try:
        # Log the summary text right before it is written to the file
        logging.info(f"Saving Summary (first 500 characters): {summary[:500]}...")

        if file_path.endswith('.txt'):
            with open(file_path, 'w', encoding='utf-8') as file:
                file.write(summary)
        elif file_path.endswith('.docx'):
            doc = Document()
            doc.add_paragraph(summary)
            doc.save(file_path)
        elif file_path.endswith('.pdf'):
            pdf = SimpleDocTemplate(file_path, pagesize=letter)
            styles = getSampleStyleSheet()
            summary_paragraph = Paragraph(summary, styles['Normal'])
            pdf.build([summary_paragraph])
        else:
            raise ValueError("Unsupported file format selected.")
    except Exception as e:
        logging.error(f"An error occurred while saving the file to {file_path}: {e}")
        raise RuntimeError(f"An error occurred while saving the file: {e}")

    logging.info(f"Summary saved successfully to {file_path}")
